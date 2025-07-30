from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
import google.generativeai as genai
import logging
import os
import io

# Initialize Flask app
app = Flask(__name__)
CORS(app)
logging.basicConfig(level=logging.INFO)

# Load API Key from .env or Render Environment
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY")

if not API_KEY:
    raise ValueError("❌ GEMINI_API_KEY is missing. Please set it in .env or Render dashboard.")

# Configure Gemini
genai.configure(api_key=API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/generate", methods=["POST"])
def generate_quiz():
    try:
        data = request.get_json()
        topic = data.get("topic", "").strip()
        num_questions = int(data.get("num_questions", 5))
        difficulty = data.get("difficulty", "easy").lower()

        # Input validation
        if not topic:
            return jsonify({"error": "Topic is required"}), 400
        if not (1 <= num_questions <= 20):
            return jsonify({"error": "Choose between 1 and 20 questions"}), 400
        if difficulty not in ["easy", "medium", "hard"]:
            return jsonify({"error": "Difficulty must be easy, medium, or hard"}), 400

        # Gemini prompt
        prompt = f"""
        Generate {num_questions} multiple-choice questions on the topic: "{topic}".
        Difficulty level: {difficulty}.

        Each question must have 4 options (A–D) and clearly indicate the correct answer.

        Format:
        1. Question?
           A) Option A
           B) Option B
           C) Option C
           D) Option D
           ✅ Answer: Option C
        """

        response = model.generate_content(prompt)
        output = getattr(response, "text", None)

        if output:
            return jsonify({"quiz": output})
        else:
            return jsonify({"error": "Empty response from Gemini"}), 500

    except Exception as e:
        logging.exception("❌ Error generating quiz:")
        return jsonify({"error": str(e)}), 500

@app.route("/download-doc", methods=["POST"])
def download_quiz_doc():
    try:
        data = request.get_json()
        quiz_text = data.get("quiz", "").strip()

        if not quiz_text:
            return jsonify({"error": "Quiz content is missing"}), 400

        # Create Word document
        doc = Document()
        doc.add_heading("Generated Quiz", level=1)
        for line in quiz_text.split("\n"):
            doc.add_paragraph(line)

        # Write to BytesIO for download
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return send_file(
            doc_io,
            as_attachment=True,
            download_name="quiz.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logging.exception("❌ Error generating DOCX file:")
        return jsonify({"error": str(e)}), 500

@app.errorhandler(500)
def handle_server_error(error):
    return jsonify({"error": "Something went wrong on the server"}), 500

if __name__ == "__main__":
    app.run(debug=True)
